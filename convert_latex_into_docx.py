# -*- coding: utf-8 -*-
"""
Created on Sat Mar  2 18:11:45 2019

@author: yiyuezhuo

Require:
    python-docx
    https://github.com/python-openxml/python-docx/issues/320
    https://python-docx.readthedocs.io/en/latest/user/quickstart.html
    
    latex2mathml
    https://pypi.org/project/latex2mathml/
    
"""

from docx import Document
from lxml import etree

from latex2mathml.converter import convert as _latex_to_mathml
'''
mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mn>1</mn><mn>2</mn></mfrac></math>'  # From the first post
'''
mml2omml_stylesheet_path = 'MML2OMML.XSL'
xslt = etree.parse(mml2omml_stylesheet_path)
transform = etree.XSLT(xslt)

def latex_to_mathml(latex):
    '''
    Convert latex to MathML:
        latex_to_mathml(r'\sin^2(x)+\cos^2(x) \neq 1')
        
        ->
        
        b'<math xmlns="http://www.w3.org/1998/Math/MathML"><mrow><msup>
        <mi>sin</mi><mn>2</mn></msup><mo>(</mo><mi>x</mi><mo>)
        </mo><mo>+</mo><msup><mi>cos</mi><mn>2</mn></msup><mo>
        (</mo><mi>x</mi><mo>)</mo><mo>&#8800;</mo><mn>1</mn></mrow></math>'
    '''
    mathml = _latex_to_mathml(latex)
    tree = etree.fromstring(mathml)
    tree.attrib['xmlns'] = "http://www.w3.org/1998/Math/MathML"
    return etree.tostring(tree)

def mathml_to_docx(mathml, docx_path):
    '''
    Create a new docx file and insert MathMl equation into it.
    '''
    tree = etree.fromstring(mathml)
    
    # Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet
    new_dom = transform(tree)
    
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph._element.append(new_dom.getroot())
    doc.save(docx_path)

def mathml_list_to_docx_as_table(mathml_list, docx_path):
    '''
    Create a new docx file and insert MathML equations list into it as 
    nx1 table.
    '''
    doc = Document()
    #paragraph = doc.add_paragraph()
    table = doc.add_table(rows=len(mathml_list), cols=1)
    for i, mathml in enumerate(mathml_list):
        #print(mathml)
        tree = etree.fromstring(mathml)
        new_dom = transform(tree)
        
        paragraph = table.cell(i,0).add_paragraph()
        paragraph._element.append(new_dom.getroot())
        #table.cell(i,0).paragraphs[0]._element.append(new_dom.getroot())
    
    doc.save(docx_path)

def mathml_list_to_docx_as_paragraph(mathml_list, docx_path):
    '''
    Create a new docx file and insert MathML equations list into it as paragraph.
    '''
    doc = Document()
    
    #table = doc.add_table(rows=len(mathml_list), cols=1)
    for i,mathml in enumerate(mathml_list):
        tree = etree.fromstring(mathml)
        new_dom = transform(tree)
        
        paragraph = doc.add_paragraph()
        paragraph._element.append(new_dom.getroot())
    
    doc.save(docx_path)

class LaTeX(str):
    '''
    generate_table will use it to identify equation.
    '''
    pass

def generate_table(mat, docx_path):
    '''
    Generate a table including cool equations in docx file.
    '''
    m = len(mat)
    n = len(mat[0])
    
    doc = Document()
    table = doc.add_table(rows=m, cols=n)
    for i in range(m):
        for j in range(n):
            content = mat[i][j]
            if isinstance(content, LaTeX):
                tree = etree.fromstring(latex_to_mathml(content))
                new_dom = transform(tree)
                
                paragraph = table.cell(i,j).add_paragraph()
                paragraph._element.append(new_dom.getroot())
            else:
                content = str(mat[i][j])
                table.cell(i, j).text = content
                
    doc.save(docx_path)
